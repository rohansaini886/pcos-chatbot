import os
from flask import Flask, request, render_template, jsonify, session, send_from_directory
from dotenv import load_dotenv
from PIL import Image
import google.generativeai as genai
from docx import Document

# ─── CONFIG ──────────────────────────────────────────────────────────────────
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "replace-with-your-own-secret")

UPLOAD_FOLDER = "static/uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# multimodal-capable model
model = genai.GenerativeModel("gemini-1.5-flash")


# ─── HELPERS ─────────────────────────────────────────────────────────────────
def call_gemini(segments):
    resp = model.generate_content(segments, stream=False)
    return resp.text

def classify_intent(text: str) -> str:
    prompt = (
        "You are an intent classifier for a PCOS chatbot. "
        "Classify the user text below into exactly one label: "
        "`greeting`, `lifestyle_suggestions`, `general_pcos_qa`, or `fallback`.\n\n"
        f"User text: \"{text}\"\n\n"
        "Reply with only the label."
    )
    return model.generate_content([prompt], stream=False).text.strip().lower()


# ─── ROUTES ──────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    session.clear()
    return render_template("index.html")

@app.route('/download/<filename>')
def download_file(filename):
    # Serve the generated DOCX file for download
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

@app.route('/analyze', methods=['POST'])
def analyze():
    # initialize session history
    if 'history' not in session:
        session['history'] = []

    desc = (request.form.get('description') or "").strip()
    file = request.files.get('file')
    user_content = desc if desc else "[image upload]"
    session['history'].append({'role': 'user', 'content': user_content})

    # build the turn-by-turn chat for context (not strictly needed for image branch)
    system_prompt = (
        "You are a helpful assistant specialized in Polycystic Ovary Syndrome (PCOS). "
        "You can analyze ovarian MRI images, answer PCOS-related or general queries, and maintain context."
    )
    convo = [system_prompt]
    for turn in session['history']:
        prefix = "User:" if turn['role'] == 'user' else "Assistant:"
        convo.append(f"{prefix} {turn['content']}")
    convo.append("Assistant:")

    # ─── IMAGE UPLOAD BRANCH ────────────────────────────────────────────────────
    if file and file.filename:
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(filepath)
        img = Image.open(filepath)

        # 1️⃣ Check for PCOS features
        check_prompt = (
            "You are a radiologist. Examine the following pelvic MRI image. "
            "Answer only 'Yes' if it shows hallmarks of PCOS "
            "(e.g., multiple peripheral follicles, enlarged ovarian volume), otherwise answer 'No'."
        )
        check = call_gemini([check_prompt, img]).strip().lower()
        if not check.startswith("yes"):
            return jsonify({
                'error': "Uploaded MRI does not display PCOS features, unable to generate report."
            }), 400

        # 2️⃣ Generate full PCOS MRI report
        report_prompt = (
            "Based on confirmation that this MRI shows PCOS, produce a structured radiology report "
            "with sections: Radiological Observations, Clinical Interpretation, "
            "Lifestyle/Wellness Suggestions, and Specialist Referral Guidance."
        )
        report = call_gemini([report_prompt, img])
        session['history'].append({'role':'assistant','content':report})

        # 3️⃣ Save as DOCX
        doc = Document()
        doc.add_heading('PCOS MRI Analysis Report', level=1)
        for line in report.split('\n'):
            doc.add_paragraph(line)
        base, _ = os.path.splitext(file.filename)
        doc_filename = f"{base}_report.docx"
        doc_path = os.path.join(app.config['UPLOAD_FOLDER'], doc_filename)
        doc.save(doc_path)

        return jsonify({
            'result': report,
            'image_url': f"/{filepath.replace(os.path.sep, '/')}",
            'download_url': f"/download/{doc_filename}"
        }), 200

    # ─── TEXT-ONLY BRANCH ────────────────────────────────────────────────────────
    intent = classify_intent(desc)

    # Greeting
    if intent == "greeting" or desc.lower() in {"hello","hi","hey"}:
        reply = "Hello! 👋 I'm your PCOS assistant—how can I help you today?"
    # Lifestyle / wellness suggestions
    elif intent == "lifestyle_suggestions":
        prompt = (
            "You are an expert in PCOS management. Provide only lifestyle and wellness suggestions "
            f"for this request: \"{desc}\". Focus exclusively on exercise routines and dietary guidelines."
        )
        reply = call_gemini([prompt])
    # General PCOS Q&A
    elif intent == "general_pcos_qa":
        prompt = (
            "You are an expert in Polycystic Ovary Syndrome (PCOS). "
            f"Answer this question clearly and in a structured way: \"{desc}\"."
        )
        reply = call_gemini([prompt])
    # Fallback: concise answer to any other question
    else:
        prompt = (
            "Answer the following question concisely and accurately:\n\n"
            f"\"{desc}\""
        )
        reply = call_gemini([prompt])

    session['history'].append({'role':'assistant','content':reply})
    return jsonify({'result': reply}), 200

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))  
    app.run(host="0.0.0.0", port=port)
